from docx import Document
from docx.shared import Pt

# 用户可以在此设置字体和首行缩进
FONT_NAME = 'FangSong'  # 设置默认字体为仿宋
FONT_SIZE = Pt(12)  # 默认字体大小为 12pt
INDENT_FIRST_LINE = True  # 是否首行缩进，默认为True

# 加载文字来源文档
source_doc = Document("source.docx")

# 加载稿纸模板
template_docs = [Document("template.docx")]

# 提取源文档的段落及文字信息
source_paragraphs = []
for paragraph in source_doc.paragraphs:
    paragraph_runs = []
    for run in paragraph.runs:
        for char in run.text:
            paragraph_runs.append({
                "char": char,
                "bold": run.bold,
                "italic": run.italic,
                "font_name": run.font.name,
                "font_size": run.font.size,
            })
    if paragraph_runs:  # 跳过空段落
        source_paragraphs.append(paragraph_runs)

# 初始化索引和状态
para_index = 0  # 当前段落索引
char_index = 0  # 当前段落内字符索引
current_row = 0  # 当前行索引
current_col = 0  # 当前列索引

# 遍历模板表格逐字插入
for doc_idx, doc in enumerate(template_docs):
    table = doc.tables[0]
    rows, cols = len(table.rows), len(table.columns)  # 表格的行列数

    while para_index < len(source_paragraphs):  # 遍历段落
        paragraph = source_paragraphs[para_index]
        while char_index < len(paragraph):  # 遍历段落中的字符
            # 首行缩进：段落第一个字符前空两格
            if char_index == 0 and current_col <= 2 and INDENT_FIRST_LINE:
                current_col = 2  # 将第一个字符之前留空两格

            # 获取字符的格式信息
            char_info = paragraph[char_index]

            # 插入字符到单元格
            cell = table.cell(current_row, current_col)
            cell.text = char_info["char"]  # 填入字符
            run = cell.paragraphs[0].runs[0]

            # 检测和应用字符的字体格式
            if char_info["bold"] is not None:
                run.bold = char_info["bold"]
            if char_info["italic"] is not None:
                run.italic = char_info["italic"]

            # 设置字体名称为用户定义的字体，默认是仿宋
            font_name = char_info["font_name"] if char_info["font_name"] else FONT_NAME
            run.font.name = font_name

            # 设置字体大小（如果字体大小为空，设置为默认 12pt）
            if char_info["font_size"]:
                run.font.size = char_info["font_size"]
            else:
                run.font.size = FONT_SIZE  # 默认字体大小为 12pt

            # 移动到下一个单元格
            char_index += 1
            current_col += 1

            # 如果列满，则换行
            if current_col >= cols:
                current_col = 0
                current_row += 1

            # 如果表格写满，切换到下一个模板
            if current_row >= rows:
                doc.save(f"output_part{doc_idx + 1}.docx")
                print(f"保存为 output_part{doc_idx + 1}.docx")
                current_row = 0
                current_col = 0
                break

        # 当前段落写完，切换到下一个段落
        if char_index >= len(paragraph):
            para_index += 1
            char_index = 0
            current_row += 1
            current_col = 0  # 下一段换行开始写

    # 保存当前模板
    doc.save(f"output_part{doc_idx + 1}.docx")
    print(f"保存为 output_part{doc_idx + 1}.docx")

    # 如果所有段落处理完成，则停止
    if para_index >= len(source_paragraphs):
        break
